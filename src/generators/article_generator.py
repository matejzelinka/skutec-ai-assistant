from collections import OrderedDict
from pathlib import Path

from docx import Document

from src.utils.formatter import (
    parse_datetime,
    format_date,
    format_time,
    format_weekday,
    location_sentence,
)


def generate_article(event: dict, output_path: Path) -> None:
    start = parse_datetime(event["from"])
    end = parse_datetime(event["to"])

    parts = []
    streets = []

    for location in event["locations"]:
        part = location["part"]
        street = location["street"]

        if part not in parts:
            parts.append(part)

        if (
            part == "Skuteč"
            and street
            and street not in streets
        ):
            streets.append(street)

    title = (
        f"Přerušení dodávky elektrické energie - "
        f"{format_date(start)}"
    )

    perex = (
        f"{format_weekday(start).capitalize()} "
        f"{format_date(start)} "
        f"od {format_time(start)} do {format_time(end)} "
        f"bude přerušena dodávka elektrické energie "
        f"{location_sentence(parts, streets)}. "
        f"Přesný rozpis naleznete níže."
    )

    sections = OrderedDict()

    for location in event["locations"]:
        part = location["part"]
        street = location["street"] or ""

        sections.setdefault(part, OrderedDict())

        sections[part].setdefault(
            street,
            {
                "cp": [],
                "ev": [],
            },
        )

        if location["house_numbers"]:
            sections[part][street]["cp"].append(
                location["house_numbers"]
            )

        if location["ev_numbers"]:
            sections[part][street]["ev"].append(
                location["ev_numbers"]
            )

    document = Document()

    document.add_paragraph(title)
    document.add_paragraph(perex)

    for part, part_streets in sections.items():
        if part != "Skuteč":
            paragraph = document.add_paragraph()

            run = paragraph.add_run(part)
            run.bold = True

            for street, values in part_streets.items():
                if values["cp"]:
                    paragraph.add_run().add_break()
                    cp = ", ".join(values["cp"])
                    paragraph.add_run(f"č. p. {cp}")

                if values["ev"]:
                    paragraph.add_run().add_break()
                    ev = ", ".join(values["ev"])
                    paragraph.add_run(f"č. ev. {ev}")

        else:
            first_street = True

            for street, values in part_streets.items():
                paragraph = document.add_paragraph()

                if first_street:
                    run = paragraph.add_run(part)
                    run.bold = True
                    paragraph.add_run().add_break()
                    first_street = False

                if street:
                    run = paragraph.add_run(street)
                    run.bold = True

                if values["cp"]:
                    paragraph.add_run().add_break()
                    cp = ", ".join(values["cp"])
                    paragraph.add_run(f"č. p. {cp}")

                if values["ev"]:
                    paragraph.add_run().add_break()
                    ev = ", ".join(values["ev"])
                    paragraph.add_run(f"č. ev. {ev}")

    document.save(output_path)